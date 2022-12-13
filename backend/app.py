from fastapi import FastAPI, Request
from fastapi.responses import FileResponse
from starlette.middleware.cors import CORSMiddleware

import modal

web_app = FastAPI()
stub = modal.Stub("doc-draft")
image = modal.Image.debian_slim().pip_install(["python-docx"])
volume = "document-store"

DOC_DIR = "/docs"

web_app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)


@web_app.post("/generate")
async def generate_docx(request: Request):
    import docx

    body = await request.json()
    templates = body["clauses"]
    doc_type = body["doc"]


    if doc_type == "rofr":
        doc = docx.Document(f"{DOC_DIR}/NVCA-2020-Right-of-First-Refusal-Template.docx")
    elif doc_type == "voting":
        doc = docx.Document(f"{DOC_DIR}/NVCA-2020-Voting-Agreement.docx")
    else:
        return {"message": "success"}


    for template in templates:
        if template["doc"] != doc_type:
            continue
        for _, config in template["template"].items():
            value = config["value"]
            for position in config["positions"]:
                paragraph = position["paragraph"]
                run = position["run"]
                text = doc.paragraphs[paragraph].runs[run].text
                start = position["runStartStr"]
                end = position["runEndStr"]
                start_index = text.find(start) if start else 0
                end_index = text.find(end) if end else len(text)
                if start is None and value is not None and type(value) is str:
                    doc.paragraphs[paragraph].runs[run].text = value + text[end_index:len(text)]
                elif value and type(value) is str:
                    doc.paragraphs[paragraph].runs[run].text = text[0:start_index + len(start) + 1] + value + text[end_index:len(text)]
        doc.save(f"{DOC_DIR}/{doc_type}.docx")
    return {"message": "success"}


@web_app.get("/download")
def download(doc: str):
    if doc == "voting":
        doc = "NVCA_Voting_Agreement_Demo.docx"
        return FileResponse(
            f"{DOC_DIR}/{doc}",
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=NVCA_Voting_Agreement.docx"},
        )
    return FileResponse(
        f"{DOC_DIR}/{doc}.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={doc}.docx"},
    )


@stub.asgi(
    image=image,
    shared_volumes={DOC_DIR: modal.Secret.from_name(volume)}
)
def fastapi_app():
    return web_app